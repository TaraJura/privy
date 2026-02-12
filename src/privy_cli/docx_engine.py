from __future__ import annotations

from dataclasses import dataclass
from typing import Iterable

from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .types import SpanReplacement

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_R = f"{{{_W_NS}}}r"
_W_HYPERLINK = f"{{{_W_NS}}}hyperlink"


@dataclass(frozen=True)
class ParagraphRef:
    paragraph: Paragraph
    location: str


def _all_runs(paragraph: Paragraph) -> list[Run]:
    """Return all Run objects in document order, including those inside hyperlinks."""
    runs: list[Run] = []
    for child in paragraph._element:
        if child.tag == _W_R:
            runs.append(Run(child, paragraph))
        elif child.tag == _W_HYPERLINK:
            for r_elem in child.findall(_W_R):
                runs.append(Run(r_elem, paragraph))
    return runs


def iter_document_paragraphs(doc: DocxDocument) -> Iterable[ParagraphRef]:
    for idx, paragraph in enumerate(doc.paragraphs):
        yield ParagraphRef(paragraph=paragraph, location=f"body:{idx}")

    for t_idx, table in enumerate(doc.tables):
        yield from _iter_table_paragraphs(table, prefix=f"body-table:{t_idx}")

    for s_idx, section in enumerate(doc.sections):
        for idx, paragraph in enumerate(section.header.paragraphs):
            yield ParagraphRef(paragraph=paragraph, location=f"section:{s_idx}:header:{idx}")
        for t_idx, table in enumerate(section.header.tables):
            yield from _iter_table_paragraphs(table, prefix=f"section:{s_idx}:header-table:{t_idx}")

        for idx, paragraph in enumerate(section.footer.paragraphs):
            yield ParagraphRef(paragraph=paragraph, location=f"section:{s_idx}:footer:{idx}")
        for t_idx, table in enumerate(section.footer.tables):
            yield from _iter_table_paragraphs(table, prefix=f"section:{s_idx}:footer-table:{t_idx}")


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in _all_runs(paragraph))


def apply_replacements_to_paragraph(paragraph: Paragraph, replacements: list[SpanReplacement]) -> int:
    if not replacements:
        return 0

    runs = _all_runs(paragraph)
    run_bounds = []
    cursor = 0
    for run in runs:
        text = run.text
        length = len(text)
        run_bounds.append((run, cursor, cursor + length))
        cursor += length

    if not run_bounds:
        return 0

    changed = 0
    for replacement in sorted(replacements, key=lambda r: (r.start, r.end), reverse=True):
        if replacement.end <= replacement.start:
            continue

        overlaps: list[tuple[object, int, int]] = []
        for run, run_start, run_end in run_bounds:
            if run_start < replacement.end and run_end > replacement.start:
                overlaps.append((run, run_start, run_end))

        if not overlaps:
            continue

        first = True
        for run, run_start, run_end in overlaps:
            run_text = run.text
            local_start = max(replacement.start, run_start) - run_start
            local_end = min(replacement.end, run_end) - run_start
            insert_text = replacement.replacement if first else ""
            new_text = run_text[:local_start] + insert_text + run_text[local_end:]
            if new_text != run_text:
                run.text = new_text
                changed += 1
            first = False

    return changed


def _iter_table_paragraphs(table: Table, prefix: str) -> Iterable[ParagraphRef]:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            yield from _iter_cell_paragraphs(cell, prefix=f"{prefix}:r{r_idx}c{c_idx}")


def _iter_cell_paragraphs(cell: _Cell, prefix: str) -> Iterable[ParagraphRef]:
    for p_idx, paragraph in enumerate(cell.paragraphs):
        yield ParagraphRef(paragraph=paragraph, location=f"{prefix}:p{p_idx}")

    for t_idx, table in enumerate(cell.tables):
        yield from _iter_table_paragraphs(table, prefix=f"{prefix}:table:{t_idx}")
