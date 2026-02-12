from pathlib import Path

from docx import Document

from privy_cli.anonymizer import anonymize_docx, deanonymize_docx
from privy_cli.detector import HeuristicDetector


def _document_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def test_anonymize_then_deanonymize_docx_keeps_text_and_styles(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    anonymized_path = tmp_path / "anonymized.docx"
    restored_path = tmp_path / "restored.docx"
    map_path = tmp_path / "mapping.json"

    source_doc = Document()
    paragraph = source_doc.add_paragraph()
    person_run = paragraph.add_run("John Doe")
    person_run.bold = True
    paragraph.add_run(" from Acme LLC lives at 123 Main Street. Reach him at john.doe@example.com.")
    source_doc.save(str(input_path))

    detector = HeuristicDetector()
    anonymize_docx(
        input_path=input_path,
        output_path=anonymized_path,
        map_path=map_path,
        detector=detector,
        entity_types=["PERSON", "COMPANY", "ADDRESS", "EMAIL"],
    )

    anonymized_doc = Document(str(anonymized_path))
    assert anonymized_doc.paragraphs[0].runs[0].bold is True
    assert "PERSON_" in anonymized_doc.paragraphs[0].text
    assert "COMPANY_" in anonymized_doc.paragraphs[0].text
    assert "ADDRESS_" in anonymized_doc.paragraphs[0].text

    deanonymize_docx(
        input_path=anonymized_path,
        output_path=restored_path,
        map_path=map_path,
    )

    assert _document_text(restored_path) == _document_text(input_path)


def test_legal_role_labels_are_not_anonymized(tmp_path: Path) -> None:
    """THE CONSULTANT, THE CLIENT etc. are legal role labels, not entities."""
    input_path = tmp_path / "input.docx"
    anonymized_path = tmp_path / "anonymized.docx"
    map_path = tmp_path / "mapping.json"

    source_doc = Document()
    source_doc.add_paragraph(
        "THE CONSULTANT: John Smith, an individual residing at 123 Main Street."
    )
    source_doc.save(str(input_path))

    detector = HeuristicDetector()
    anonymize_docx(
        input_path=input_path,
        output_path=anonymized_path,
        map_path=map_path,
        detector=detector,
        entity_types=["PERSON", "COMPANY", "ADDRESS"],
    )

    anonymized_text = _document_text(anonymized_path)
    assert "THE CONSULTANT" in anonymized_text, "Legal role label should not be anonymized"
    assert "PERSON_" in anonymized_text, "Actual person name should be anonymized"
